"""
document_export.py
──────────────────
Pure-Python document export — no LibreOffice needed.

  html_to_docx(html, filename)  →  bytes   (DOCX)
  html_to_pdf(html, title)      →  bytes   (PDF via reportlab)

The HTML comes from TipTap's editor output, so we handle:
  headings  <h1> <h2> <h3>
  paragraph <p>
  bold      <strong> <b>
  italic    <em> <i>
  underline <u>
  lists     <ul>/<ol> + <li>
  line-break <br>
"""

import io
import re
from html.parser import HTMLParser
from typing import Optional


# ── DOCX export via python-docx ───────────────────────────────────────────────

def html_to_docx(html: str, title: str = "Contract") -> bytes:
    """Convert HTML → a professionally styled DOCX document."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page margins (1 inch all around) ─────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── Base font ─────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    # ── Heading styles ────────────────────────────────────────────────────
    for style_name, size, hex_color, sp_before, sp_after in [
        ("Heading 1", 15, "1E2D4A", 20, 8),
        ("Heading 2", 13, "1E3A5F", 14, 6),
        ("Heading 3", 11, "2D4A6A", 10, 4),
    ]:
        try:
            st = doc.styles[style_name]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.bold = True
            st.font.color.rgb = RGBColor(
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16),
            )
            st.paragraph_format.space_before = Pt(sp_before)
            st.paragraph_format.space_after = Pt(sp_after)
            st.paragraph_format.keep_with_next = True
        except Exception:
            pass

    # ── Title block ───────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(8)
    run = tp.add_run(title.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Thin bottom border under title
    pPr = tp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '2')
    bot.set(qn('w:color'), '1E293B')
    pBdr.append(bot)
    pPr.append(pBdr)

    # Spacer after title
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)

    # ── Content blocks ────────────────────────────────────────────────────
    blocks = _parse_html_blocks(html)

    for block in blocks:
        btype  = block["type"]
        runs   = block["runs"]
        indent = block.get("indent", 0)

        if btype.startswith("h") and btype != "hr":
            level = int(btype[1])
            heading_style = f"Heading {min(level, 3)}"
            try:
                p = doc.add_paragraph(style=heading_style)
            except Exception:
                p = doc.add_paragraph()
            _add_runs(p, runs)

        elif btype == "li_bullet":
            p = doc.add_paragraph(style="List Bullet")
            if indent:
                p.paragraph_format.left_indent = Pt(18 * indent)
            _add_runs(p, runs)

        elif btype == "li_ordered":
            p = doc.add_paragraph(style="List Number")
            if indent:
                p.paragraph_format.left_indent = Pt(18 * indent)
            _add_runs(p, runs)

        elif btype == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr2 = p._p.get_or_add_pPr()
            pBdr2 = OxmlElement('w:pBdr')
            bot2 = OxmlElement('w:bottom')
            bot2.set(qn('w:val'), 'single')
            bot2.set(qn('w:sz'), '4')
            bot2.set(qn('w:space'), '1')
            bot2.set(qn('w:color'), 'CBD5E1')
            pBdr2.append(bot2)
            pPr2.append(pBdr2)

        elif btype == "br":
            sp2 = doc.add_paragraph()
            sp2.paragraph_format.space_before = Pt(0)
            sp2.paragraph_format.space_after = Pt(4)

        else:  # "p"
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_runs(p, runs)

    # ── Footer: "Title · Page X of Y" ────────────────────────────────────
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    FOOTER_GRAY = RGBColor(0x94, 0xA3, 0xB8)

    def _footer_run(text: str = ""):
        r = fp.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = FOOTER_GRAY
        return r

    _footer_run(f"{title}  ·  Page ")

    r_page = _footer_run()
    for fld_type, instr in [('begin', None), (None, ' PAGE '), ('end', None),
                             (None, None)]:
        if fld_type:
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), fld_type)
            r_page._r.append(fc)
        elif instr:
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = instr
            r_page._r.append(it)

    _footer_run(" of ")

    r_total = _footer_run()
    for fld_type, instr in [('begin', None), (None, ' NUMPAGES '), ('end', None)]:
        if fld_type:
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), fld_type)
            r_total._r.append(fc)
        elif instr:
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = instr
            r_total._r.append(it)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_runs(paragraph, runs: list):
    for r in runs:
        text = r.get("text", "")
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold      = r.get("bold",      False)
        run.italic    = r.get("italic",    False)
        run.underline = r.get("underline", False)


# ── PDF export via reportlab ──────────────────────────────────────────────────

def html_to_pdf(html: str, title: str = "Contract") -> bytes:
    """Convert TipTap HTML → PDF bytes using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.platypus import ListFlowable, ListItem

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles   = getSampleStyleSheet()
    DARK     = colors.HexColor("#1e293b")
    BLUE     = colors.HexColor("#3730a3")
    GRAY     = colors.HexColor("#475569")

    title_style = ParagraphStyle(
        "ContractTitle",
        parent=styles["Heading1"],
        fontSize=18,
        textColor=DARK,
        spaceAfter=20,
        alignment=1,   # CENTER
    )
    h1_style = ParagraphStyle("CH1", parent=styles["Heading1"], fontSize=15, textColor=BLUE, spaceBefore=14, spaceAfter=6)
    h2_style = ParagraphStyle("CH2", parent=styles["Heading2"], fontSize=13, textColor=BLUE, spaceBefore=10, spaceAfter=4)
    h3_style = ParagraphStyle("CH3", parent=styles["Heading3"], fontSize=11, textColor=DARK, spaceBefore=8,  spaceAfter=3)
    body_style = ParagraphStyle(
        "CBody",
        parent=styles["Normal"],
        fontSize=10.5,
        textColor=GRAY,
        leading=16,
        spaceAfter=6,
    )

    story = [
        Paragraph(title, title_style),
        Spacer(1, 0.3 * cm),
    ]

    blocks = _parse_html_blocks(html)
    list_buffer: list = []   # accumulate consecutive li items

    def flush_list(ordered: bool):
        nonlocal list_buffer
        if not list_buffer:
            return
        items = [ListItem(Paragraph(_runs_to_rl(r), body_style), leftIndent=20) for r in list_buffer]
        story.append(ListFlowable(items, bulletType="1" if ordered else "bullet", start="1" if ordered else None))
        story.append(Spacer(1, 0.2 * cm))
        list_buffer = []

    prev_ordered = False

    for block in blocks:
        btype = block["type"]
        is_li = btype.startswith("li_")
        ordered = btype == "li_ordered"

        if not is_li and list_buffer:
            flush_list(prev_ordered)

        if btype == "h1":
            story.append(Paragraph(_runs_to_rl(block["runs"]), h1_style))
        elif btype == "h2":
            story.append(Paragraph(_runs_to_rl(block["runs"]), h2_style))
        elif btype == "h3":
            story.append(Paragraph(_runs_to_rl(block["runs"]), h3_style))
        elif is_li:
            list_buffer.append(block["runs"])
            prev_ordered = ordered
        elif btype == "br":
            story.append(Spacer(1, 0.3 * cm))
        else:
            text = _runs_to_rl(block["runs"])
            if text.strip():
                story.append(Paragraph(text, body_style))

    if list_buffer:
        flush_list(prev_ordered)

    doc.build(story)
    return buf.getvalue()


def _runs_to_rl(runs: list) -> str:
    """Convert run dicts to ReportLab XML markup."""
    parts = []
    for r in runs:
        text = r.get("text", "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if not text:
            continue
        if r.get("bold"):
            text = f"<b>{text}</b>"
        if r.get("italic"):
            text = f"<i>{text}</i>"
        if r.get("underline"):
            text = f"<u>{text}</u>"
        parts.append(text)
    return "".join(parts) or " "


# ── HTML parser ───────────────────────────────────────────────────────────────

class _BlockParser(HTMLParser):
    """
    Converts TipTap HTML into a flat list of block dicts:
      { type: "p"|"h1"|"h2"|"h3"|"li_bullet"|"li_ordered"|"br",
        runs: [{text, bold, italic, underline}] }
    """

    BLOCK_TAGS  = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "br", "div", "hr"}
    INLINE_BOLD = {"strong", "b"}
    INLINE_ITAL = {"em", "i"}
    INLINE_UND  = {"u"}

    def __init__(self):
        super().__init__()
        self.blocks:     list  = []
        self._cur_block: Optional[dict] = None
        self._in_ol:     int   = 0
        self._bold:      int   = 0
        self._italic:    int   = 0
        self._underline: int   = 0

    def _start_block(self, btype: str):
        self._flush()
        self._cur_block = {"type": btype, "runs": []}

    def _flush(self):
        if self._cur_block and (
            self._cur_block["runs"]
            or self._cur_block["type"] == "br"
        ):
            self.blocks.append(self._cur_block)
        self._cur_block = None

    def _add_text(self, text: str):
        if self._cur_block is None:
            self._cur_block = {"type": "p", "runs": []}
        if not text:
            return
        self._cur_block["runs"].append({
            "text":      text,
            "bold":      self._bold > 0,
            "italic":    self._italic > 0,
            "underline": self._underline > 0,
        })

    # ── tag handlers ─────────────────────────────────────────────────────

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._start_block(tag[:2])        # h1/h2/h3
        elif tag == "p":
            self._start_block("p")
        elif tag == "div":
            if self._cur_block is None:
                self._start_block("p")
        elif tag == "br":
            self._flush()
            self.blocks.append({"type": "br", "runs": []})
        elif tag == "hr":
            self._flush()
            self.blocks.append({"type": "hr", "runs": []})
        elif tag == "ol":
            self._in_ol += 1
        elif tag == "ul":
            pass   # ul depth tracked implicitly
        elif tag == "li":
            btype = "li_ordered" if self._in_ol > 0 else "li_bullet"
            self._start_block(btype)
        elif tag in self.INLINE_BOLD:
            self._bold += 1
        elif tag in self.INLINE_ITAL:
            self._italic += 1
        elif tag in self.INLINE_UND:
            self._underline += 1

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "div"):
            self._flush()
        elif tag == "ol":
            self._in_ol = max(0, self._in_ol - 1)
        elif tag in self.INLINE_BOLD:
            self._bold = max(0, self._bold - 1)
        elif tag in self.INLINE_ITAL:
            self._italic = max(0, self._italic - 1)
        elif tag in self.INLINE_UND:
            self._underline = max(0, self._underline - 1)

    def handle_data(self, data):
        self._add_text(data)

    def handle_entityref(self, name):
        entities = {"amp": "&", "lt": "<", "gt": ">", "nbsp": " ",
                    "quot": '"', "apos": "'"}
        self._add_text(entities.get(name, ""))

    def handle_charref(self, name):
        try:
            ch = chr(int(name[1:], 16) if name.startswith("x") else int(name))
            self._add_text(ch)
        except Exception:
            pass


def _parse_html_blocks(html: str) -> list:
    """Parse HTML string into a list of block dicts."""
    # Normalise self-closing <br/> → <br>
    html = re.sub(r"<br\s*/?>", "<br>", html, flags=re.IGNORECASE)
    parser = _BlockParser()
    parser.feed(html)
    parser._flush()
    return parser.blocks
