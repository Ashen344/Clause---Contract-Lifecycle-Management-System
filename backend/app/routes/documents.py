import os
import shutil
import re
import uuid
import httpx
import xml.etree.ElementTree as ET
from datetime import datetime
from fastapi import APIRouter, HTTPException, Query, UploadFile, File, Form, Depends, Request
from fastapi.responses import FileResponse
from bson import ObjectId
from pydantic import BaseModel
from app.config import (
    contracts_collection,
    UPLOAD_DIR,
    ALLOWED_EXTENSIONS,
    MAX_FILE_SIZE,
    COLLABORA_INTERNAL_URL,
    WOPI_BASE_URL,
)
from app.middleware.auth import get_current_user, get_optional_user
from app.routes.wopi import make_wopi_token

MIME_MAP = {
    ".pdf": "application/pdf",
    ".txt": "text/plain; charset=utf-8",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".rtf": "application/rtf",
    ".odt": "application/vnd.oasis.opendocument.text",
}

router = APIRouter(prefix="/api/documents", tags=["Documents"])


@router.post("/upload/{contract_id}")
async def upload_document(
    contract_id: str,
    file: UploadFile = File(...),
    change_notes: str = Form(default=""),
    current_user: dict = Depends(get_optional_user),
):
    """Upload a document to a contract."""
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    # Validate file extension
    _, ext = os.path.splitext(file.filename or "")
    ext = ext.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{ext}' not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Read file content
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File size exceeds 20MB limit")

    # Save file to disk
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    file_id = uuid.uuid4().hex
    stored_filename = f"{file_id}{ext}"
    file_path = os.path.join(UPLOAD_DIR, stored_filename)

    with open(file_path, "wb") as f:
        f.write(content)

    # Build version entry
    current_version = contract.get("current_version", 0)
    new_version = current_version + 1
    user_id = current_user["user_id"] if current_user else "unknown"

    version_entry = {
        "version_number": new_version,
        "file_url": stored_filename,
        "original_filename": file.filename,
        "file_size": len(content),
        "file_type": ext,
        "uploaded_by": user_id,
        "uploaded_at": datetime.utcnow(),
        "change_notes": change_notes or None,
    }

    # Update contract in DB
    contracts_collection.update_one(
        {"_id": ObjectId(contract_id)},
        {
            "$push": {"versions": version_entry},
            "$set": {
                "file_url": stored_filename,
                "current_version": new_version,
                "updated_at": datetime.utcnow(),
            },
        },
    )

    return {
        "message": "Document uploaded successfully",
        "version": new_version,
        "filename": file.filename,
        "file_size": len(content),
        "file_type": ext,
    }


@router.get("/download/{contract_id}")
async def download_document(contract_id: str, version: int = 0):
    """Download a document. If version=0 (default), downloads the latest."""
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    versions = contract.get("versions", [])
    if not versions:
        raise HTTPException(status_code=404, detail="No documents uploaded for this contract")

    # Find the requested version
    if version > 0:
        target = next((v for v in versions if v["version_number"] == version), None)
        if not target:
            raise HTTPException(status_code=404, detail=f"Version {version} not found")
    else:
        target = versions[-1]  # latest

    file_path = os.path.join(UPLOAD_DIR, target["file_url"])
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found on disk")

    original_name = target.get("original_filename", target["file_url"])
    return FileResponse(
        path=file_path,
        filename=original_name,
        media_type="application/octet-stream",
    )


@router.get("/view/{contract_id}")
async def view_document(contract_id: str, version: int = 0):
    """Serve a document inline so the browser can render it (PDF viewer, text, etc.)."""
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    versions = contract.get("versions", [])
    if not versions:
        raise HTTPException(status_code=404, detail="No documents uploaded for this contract")

    if version > 0:
        target = next((v for v in versions if v["version_number"] == version), None)
        if not target:
            raise HTTPException(status_code=404, detail=f"Version {version} not found")
    else:
        target = versions[-1]

    file_path = os.path.join(UPLOAD_DIR, target["file_url"])
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found on disk")

    ext = target.get("file_type", ".pdf")
    media_type = MIME_MAP.get(ext, "application/octet-stream")
    original_name = target.get("original_filename", target["file_url"])

    return FileResponse(
        path=file_path,
        media_type=media_type,
        headers={"Content-Disposition": f'inline; filename="{original_name}"'},
    )


@router.get("/list/{contract_id}")
async def list_documents(contract_id: str):
    """List all document versions for a contract."""
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    versions = contract.get("versions", [])
    return {
        "contract_id": contract_id,
        "current_version": contract.get("current_version", 0),
        "documents": [
            {
                "version_number": v.get("version_number"),
                "original_filename": v.get("original_filename", v.get("file_url", "")),
                "file_size": v.get("file_size"),
                "file_type": v.get("file_type"),
                "uploaded_by": v.get("uploaded_by"),
                "uploaded_at": v.get("uploaded_at"),
                "change_notes": v.get("change_notes"),
            }
            for v in versions
        ],
    }


class SaveTextRequest(BaseModel):
    text: str


@router.get("/text/{contract_id}")
async def get_document_text(contract_id: str):
    """Return the editable text content stored for a contract document.
    Falls back to live file extraction when no text is cached yet."""
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    # 1. Use already-extracted text if available
    stored_text = contract.get("extracted_text", "")
    versions = contract.get("versions", [])
    file_type = ".txt"
    has_file = bool(versions)

    if versions:
        latest = versions[-1]
        file_type = latest.get("file_type", ".txt")

    if stored_text:
        return {"text": stored_text, "file_type": file_type, "has_file": has_file}

    # 2. Try to extract from disk
    if not versions:
        return {"text": "", "file_type": file_type, "has_file": False}

    latest = versions[-1]
    file_path = os.path.join(UPLOAD_DIR, latest.get("file_url", ""))
    ext = latest.get("file_type", ".txt")

    if not os.path.exists(file_path):
        return {"text": "", "file_type": ext, "has_file": False}

    extracted = ""
    try:
        if ext == ".pdf":
            import io
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(open(file_path, "rb").read()))
                extracted = "\n\n".join(
                    p.extract_text() or "" for p in reader.pages
                ).strip()
            except Exception:
                pass

        elif ext in (".docx", ".doc"):
            # Extract text from DOCX using python-docx
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
                # Also pull text out of tables
                for table in doc.tables:
                    for row in table.rows:
                        row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if row_cells:
                            parts.append("  |  ".join(row_cells))
                extracted = "\n\n".join(parts)
            except Exception as docx_err:
                extracted = ""

        elif ext in (".txt", ".rtf"):
            with open(file_path, "r", errors="replace") as f:
                extracted = f.read()

        elif ext == ".odt":
            # Basic ODT extraction via zipfile (ODT is a ZIP with content.xml)
            try:
                import zipfile
                with zipfile.ZipFile(file_path) as z:
                    xml = z.read("content.xml").decode("utf-8", errors="replace")
                extracted = re.sub(r"<[^>]+>", " ", xml)
                extracted = " ".join(extracted.split())
            except Exception:
                extracted = ""

    except Exception:
        extracted = ""

    # Cache for next time
    if extracted:
        contracts_collection.update_one(
            {"_id": ObjectId(contract_id)},
            {"$set": {"extracted_text": extracted, "updated_at": datetime.utcnow()}},
        )

    return {"text": extracted, "file_type": ext, "has_file": True}


@router.put("/text/{contract_id}")
async def save_document_text(
    contract_id: str,
    body: SaveTextRequest,
    current_user: dict = Depends(get_current_user),
):
    """
    Save edited document text.

    1. Writes the text to a new .txt file on disk so the preview reflects it.
    2. Adds the .txt file as a new version entry in the contract.
    3. Stores it in extracted_text for fast retrieval.
    """
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    # ── Write the edited text to a new .txt file on disk ──────────────────
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    file_id         = uuid.uuid4().hex
    stored_filename = f"{file_id}.txt"
    file_path       = os.path.join(UPLOAD_DIR, stored_filename)
    encoded         = body.text.encode("utf-8")

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(body.text)

    # ── Build a new version entry ──────────────────────────────────────────
    current_version = contract.get("current_version", 0)
    new_version     = current_version + 1
    user_id         = current_user.get("user_id", "unknown")

    # Derive a sensible original filename (e.g. "mycontract_edited.txt")
    versions = contract.get("versions", [])
    original_stem = "document"
    if versions:
        prev_name = versions[-1].get("original_filename", "document.txt")
        original_stem = os.path.splitext(prev_name)[0]
    original_filename = f"{original_stem}_edited.txt"

    version_entry = {
        "version_number":   new_version,
        "file_url":         stored_filename,
        "original_filename": original_filename,
        "file_size":        len(encoded),
        "file_type":        ".txt",
        "uploaded_by":      user_id,
        "uploaded_at":      datetime.utcnow(),
        "change_notes":     "Edited via document editor",
    }

    # ── Persist to MongoDB ─────────────────────────────────────────────────
    contracts_collection.update_one(
        {"_id": ObjectId(contract_id)},
        {
            "$push": {"versions": version_entry},
            "$set":  {
                "file_url":        stored_filename,
                "current_version": new_version,
                "extracted_text":  body.text,
                "updated_at":      datetime.utcnow(),
            },
        },
    )

    return {
        "message":  "Document saved successfully",
        "version":  new_version,
        "file_type": ".txt",
    }


# ── Generate DOCX from template content ──────────────────────────────────────

@router.post("/generate-from-template/{contract_id}")
async def generate_document_from_template(
    contract_id: str,
    current_user: dict = Depends(get_current_user),
):
    """Generate a .docx file from the contract's linked template content and attach it."""
    from app.config import templates_collection
    from docx import Document as DocxDocument  # type: ignore[import]
    from docx.shared import Inches  # type: ignore[import]
    import io

    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    template_id = contract.get("template_id")
    if not template_id or not ObjectId.is_valid(template_id):
        raise HTTPException(status_code=400, detail="Contract has no linked template")

    template = templates_collection.find_one({"_id": ObjectId(template_id)})
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    content: str = template.get("content", "")
    title: str = contract.get("title", template.get("name", "Contract"))

    # ── Build DOCX from markdown-like content ──────────────────────────────
    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    def _add_inline_bold(para, text: str):
        """Split text on **bold** markers and add runs accordingly."""
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for i, part in enumerate(parts):
            if not part:
                continue
            run = para.add_run(part)
            if i % 2 == 1:  # odd indices are inside **...**
                run.bold = True

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading 1
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            i += 1
        # Heading 2
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            i += 1
        # Heading 3
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            i += 1
        # Table separator row — skip (markdown table dividers like |---|---|)
        elif re.match(r'^\|[-| :]+\|$', line.strip()):
            i += 1
        # Table row
        elif line.strip().startswith("|") and line.strip().endswith("|"):
            # Collect all consecutive table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                row_line = lines[i].strip()
                # Skip divider rows
                if not re.match(r'^\|[-| :]+\|$', row_line):
                    cells = [c.strip() for c in row_line.strip("|").split("|")]
                    table_rows.append(cells)
                i += 1
            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=num_cols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(table_rows):
                    for ci, cell_text in enumerate(row):
                        if ci < num_cols:
                            cell = tbl.cell(ri, ci)
                            cell.text = cell_text
                            if ri == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
        # Bullet list
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_bold(p, line[2:].strip())
            i += 1
        # Horizontal rule
        elif line.strip() in ("---", "___", "***"):
            doc.add_paragraph("_" * 60)
            i += 1
        # Empty line → paragraph break
        elif line.strip() == "":
            doc.add_paragraph("")
            i += 1
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline_bold(p, line.strip())
            i += 1

    # ── Save to bytes and write to upload dir ─────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    file_bytes = buf.getvalue()

    os.makedirs(UPLOAD_DIR, exist_ok=True)
    file_id = uuid.uuid4().hex
    stored_filename = f"{file_id}.docx"
    file_path = os.path.join(UPLOAD_DIR, stored_filename)
    with open(file_path, "wb") as f:
        f.write(file_bytes)

    # ── Attach as version 1 of the contract ───────────────────────────────
    current_version = contract.get("current_version", 0)
    new_version = current_version + 1
    safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(" ", "_") or "contract"
    original_filename = f"{safe_title}.docx"

    version_entry = {
        "version_number": new_version,
        "file_url": stored_filename,
        "original_filename": original_filename,
        "file_size": len(file_bytes),
        "file_type": ".docx",
        "uploaded_by": current_user.get("user_id", "system"),
        "uploaded_at": datetime.utcnow(),
        "change_notes": f"Generated from template: {template.get('name', '')}",
    }

    contracts_collection.update_one(
        {"_id": ObjectId(contract_id)},
        {
            "$push": {"versions": version_entry},
            "$set": {
                "file_url": stored_filename,
                "current_version": new_version,
                "updated_at": datetime.utcnow(),
            },
        },
    )

    return {
        "message": "Document generated from template",
        "version": new_version,
        "filename": original_filename,
        "file_type": ".docx",
    }


# ── Rich-text export (DOCX / PDF) — no LibreOffice needed ────────────────────

class SaveHtmlRequest(BaseModel):
    html: str
    title: str = "Contract"


@router.put("/html/{contract_id}")
async def save_document_html(
    contract_id: str,
    body: SaveHtmlRequest,
    current_user: dict = Depends(get_current_user),
):
    """Save TipTap HTML content to the contract record."""
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    result = contracts_collection.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "document_html": body.html,
            "updated_at":    datetime.utcnow(),
        }},
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Contract not found")

    return {"message": "HTML saved"}


@router.get("/html/{contract_id}")
async def get_document_html(contract_id: str):
    """Return stored TipTap HTML for a contract, or '' if none."""
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one(
        {"_id": ObjectId(contract_id)},
        {"document_html": 1, "extracted_text": 1, "title": 1},
    )
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    html = contract.get("document_html", "")

    # If no HTML stored yet, wrap extracted_text in basic paragraph tags
    if not html:
        raw = contract.get("extracted_text", "")
        if raw:
            # Wrap each paragraph in <p> tags
            html = "".join(
                f"<p>{line.strip()}</p>"
                for line in raw.split("\n")
                if line.strip()
            )

    return {"html": html, "title": contract.get("title", "Contract")}


@router.post("/export/{contract_id}")
async def export_document(
    contract_id: str,
    fmt: str = Query(..., alias="format", description="docx or pdf"),
    current_user: dict = Depends(get_current_user),
):
    """
    Export the contract's rich-text content as DOCX or PDF.
    Uses the stored TipTap HTML — no LibreOffice required.
    """
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    fmt = fmt.lower().strip(".")
    if fmt not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="format must be 'docx' or 'pdf'")

    contract = contracts_collection.find_one(
        {"_id": ObjectId(contract_id)},
        {"document_html": 1, "extracted_text": 1, "title": 1},
    )
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    html  = contract.get("document_html", "")
    title = contract.get("title", "Contract")

    # Fallback: build basic HTML from extracted_text
    if not html:
        raw = contract.get("extracted_text", "")
        html = "".join(
            f"<p>{line.strip()}</p>"
            for line in raw.split("\n")
            if line.strip()
        ) or "<p>No content.</p>"

    try:
        if fmt == "docx":
            file_bytes  = html_to_docx(html, title)
            media_type  = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename    = f"{title}.docx"
        else:
            file_bytes  = html_to_pdf(html, title)
            media_type  = "application/pdf"
            filename    = f"{title}.pdf"
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Export failed: {exc}")

    from fastapi.responses import Response
    safe_name = filename.replace('"', "'")
    return Response(
        content=file_bytes,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{safe_name}"'},
    )


@router.get("/wopi-url/{contract_id}")
async def get_wopi_url(
    contract_id: str,
    request: Request,
    current_user: dict = Depends(get_current_user),
):
    """Return the Collabora Online editor URL for the contract's latest document.

    The frontend embeds this URL in an iframe to open LibreOffice in the browser.
    Collabora calls back to /wopi/files/{contract_id} to read and save the file.
    """
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    versions = contract.get("versions", [])
    if not versions:
        raise HTTPException(status_code=404, detail="No document attached to this contract")

    latest = versions[-1]
    file_type = latest.get("file_type", ".docx").lstrip(".")
    token = make_wopi_token(contract_id)

    # The WOPI source URL must be reachable from inside the Collabora Docker container.
    # host.docker.internal resolves to the Windows host from within WSL/Docker.
    wopi_src = f"{WOPI_BASE_URL}/wopi/files/{contract_id}?access_token={token}"

    # Fetch Collabora's discovery XML to get the correct action URL for this file type.
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(f"{COLLABORA_INTERNAL_URL}/hosting/discovery")
            resp.raise_for_status()
            root = ET.fromstring(resp.text)
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail=f"Cannot reach Collabora at {COLLABORA_INTERNAL_URL}: {exc}",
        )

    # Find action URL: prefer "edit", fall back to view variants for read-only formats (PDF)
    action_url = None
    for action_name in ("edit", "view_comment", "view"):
        for action in root.findall(f".//action[@name='{action_name}']"):
            if action.get("ext", "") == file_type:
                action_url = action.get("urlsrc")
                break
        if action_url:
            break

    # Final fallback: any edit action for common editable formats
    if not action_url and file_type not in ("pdf",):
        for action in root.findall(".//action[@name='edit']"):
            if action.get("ext", "") in ("docx", "odt"):
                action_url = action.get("urlsrc")
                break

    if not action_url:
        raise HTTPException(
            status_code=422,
            detail=f"Collabora has no supported action for file type '.{file_type}'",
        )

    # Rewrite internal Docker URL (https://code:9980/...) to go through the nginx
    # /collabora/ reverse proxy so the browser can reach it.
    from urllib.parse import quote
    proto = request.headers.get("x-forwarded-proto", "https")
    host = request.headers.get("host", "localhost")
    action_url = re.sub(r"https?://[^/?]+", f"{proto}://{host}", action_url, count=1)

    # Collabora action URLs end with "?" or contain template params — append WOPISrc
    if "?" in action_url:
        editor_url = f"{action_url}WOPISrc={quote(wopi_src, safe='')}"
    else:
        editor_url = f"{action_url}?WOPISrc={quote(wopi_src, safe='')}"

    return {
        "editor_url": editor_url,
        "file_type": file_type,
        "filename": latest.get("original_filename", "document"),
    }


@router.delete("/{contract_id}/{version_number}")
async def delete_document(
    contract_id: str,
    version_number: int,
    current_user: dict = Depends(get_current_user),
):
    """Delete a specific document version."""
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    versions = contract.get("versions", [])
    target = next((v for v in versions if v["version_number"] == version_number), None)
    if not target:
        raise HTTPException(status_code=404, detail="Version not found")

    # Remove file from disk
    file_path = os.path.join(UPLOAD_DIR, target["file_url"])
    if os.path.exists(file_path):
        os.remove(file_path)

    # Remove from DB
    contracts_collection.update_one(
        {"_id": ObjectId(contract_id)},
        {"$pull": {"versions": {"version_number": version_number}}},
    )

    # Update current version if needed
    remaining = [v for v in versions if v["version_number"] != version_number]
    new_current = remaining[-1]["version_number"] if remaining else 0
    new_file_url = remaining[-1]["file_url"] if remaining else None

    contracts_collection.update_one(
        {"_id": ObjectId(contract_id)},
        {"$set": {
            "current_version": new_current,
            "file_url": new_file_url,
            "updated_at": datetime.utcnow(),
        }},
    )

    return {"message": f"Version {version_number} deleted successfully"}


# ── LibreOffice conversion ────────────────────────────────────────────────────

@router.get("/libreoffice-status")
async def libreoffice_status():
    """Check whether LibreOffice is available on this server."""
    available = is_libreoffice_available()
    return {
        "available": available,
        "supported_formats": sorted(ALLOWED_TARGETS) if available else [],
        "message": (
            "LibreOffice is installed and ready."
            if available
            else (
                "LibreOffice is not installed on this server. "
                "Install it to enable DOCX↔PDF conversion. "
                "See: https://www.libreoffice.org/"
            )
        ),
    }


@router.post("/convert/{contract_id}")
async def convert_document_endpoint(
    contract_id: str,
    target_format: str = Query(
        ...,
        description="Target format: pdf, docx, txt, odt, rtf",
    ),
    current_user: dict = Depends(get_current_user),
):
    """
    Convert a contract's latest document to a different format using LibreOffice.

    Returns the converted file as a download.  The original file is not replaced.
    Supported conversions: DOCX→PDF, PDF→DOCX, DOC→PDF, DOC→DOCX, etc.
    """
    if not ObjectId.is_valid(contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    # Verify LibreOffice is available
    if not is_libreoffice_available():
        raise HTTPException(
            status_code=503,
            detail=(
                "LibreOffice is not installed on this server. "
                "A system administrator must install it to enable document conversion."
            ),
        )

    contract = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    versions = contract.get("versions", [])
    if not versions:
        raise HTTPException(
            status_code=404, detail="No document uploaded for this contract"
        )

    latest = versions[-1]
    source_path = os.path.join(UPLOAD_DIR, latest.get("file_url", ""))

    if not os.path.exists(source_path):
        raise HTTPException(status_code=404, detail="Source file not found on disk")

    # Validate target format
    fmt = target_format.lower().lstrip(".")
    if fmt not in ALLOWED_TARGETS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format '{fmt}'. Allowed: {', '.join(sorted(ALLOWED_TARGETS))}",
        )

    # Refuse pointless same-format conversions
    source_ext = latest.get("file_type", "").lstrip(".")
    if source_ext == fmt:
        raise HTTPException(
            status_code=400,
            detail=f"Source file is already in {fmt} format.",
        )

    # Convert
    tmp_dir = None
    try:
        out_path = convert_document(source_path, fmt)
        tmp_dir = os.path.dirname(out_path)

        original_stem = os.path.splitext(
            latest.get("original_filename", "contract")
        )[0]
        download_name = f"{original_stem}.{fmt}"

        MIME_FOR_FORMAT = {
            "pdf":  "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "txt":  "text/plain; charset=utf-8",
            "odt":  "application/vnd.oasis.opendocument.text",
            "rtf":  "application/rtf",
        }
        media_type = MIME_FOR_FORMAT.get(fmt, "application/octet-stream")

        # FileResponse streams the file; we clean up the temp dir afterwards
        # using a background task.
        from starlette.background import BackgroundTask

        def _cleanup():
            shutil.rmtree(tmp_dir, ignore_errors=True)

        return FileResponse(
            path=out_path,
            filename=download_name,
            media_type=media_type,
            background=BackgroundTask(_cleanup),
        )

    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    except RuntimeError as exc:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        raise HTTPException(status_code=500, detail=str(exc))
    except Exception as exc:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        raise HTTPException(
            status_code=500,
            detail=f"Conversion error: {exc}",
        )
