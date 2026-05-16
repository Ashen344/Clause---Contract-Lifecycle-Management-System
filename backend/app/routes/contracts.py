import os
import re
import uuid
from datetime import datetime

from fastapi import APIRouter, BackgroundTasks, HTTPException, Query, UploadFile, File, Depends
from pydantic import BaseModel
from typing import Optional
from bson import ObjectId
from app.models.contract import (
    ContractCreate,
    ContractUpdate,
    ContractFilter,
    ContractType,
    ContractStatus,
    WorkflowStage,
    RiskLevel,
)
from app.services.contract_service import (
    create_contract,
    get_contract,
    get_contracts,
    update_contract,
    delete_contract,
    permanent_delete_contract,
    restore_contract,
    archive_contract,
    unarchive_contract,
    purge_expired_trash,
    update_workflow_stage,
    get_dashboard_stats,
)
from app.config import contracts_collection, UPLOAD_DIR, ALLOWED_EXTENSIONS, MAX_FILE_SIZE
from app.middleware.auth import get_current_user_with_role
from app.services.audit_service import create_audit_log
from app.models.audit_log import AuditAction

router = APIRouter(prefix="/api/contracts", tags=["Contracts"])

_MONTHS = {
    "january":1,"february":2,"march":3,"april":4,"may":5,"june":6,
    "july":7,"august":8,"september":9,"october":10,"november":11,"december":12,
    "jan":1,"feb":2,"mar":3,"apr":4,"jun":6,"jul":7,
    "aug":8,"sep":9,"oct":10,"nov":11,"dec":12,
}

def _parse_date_str(s: str):
    s = s.strip().rstrip(".,;)")
    fmts = [
        "%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y", "%d/%m/%Y", "%Y/%m/%d",
        "%d %B %Y", "%B %d, %Y", "%B %d %Y",
        "%d %b %Y", "%b %d, %Y", "%b %d %Y",
        "%B %dst %Y", "%B %dnd %Y", "%B %drd %Y", "%B %dth %Y",
    ]
    for fmt in fmts:
        try:
            return datetime.strptime(s, fmt)
        except ValueError:
            pass
    # Try "1 Jan 2025" or "Jan 1 2025" via regex
    m = re.match(r'(\d{1,2})\s+([A-Za-z]+)\s+(\d{4})', s)
    if m:
        day, mon, year = int(m.group(1)), _MONTHS.get(m.group(2).lower()), int(m.group(3))
        if mon:
            try:
                return datetime(year, mon, day)
            except ValueError:
                pass
    m = re.match(r'([A-Za-z]+)\s+(\d{1,2}),?\s+(\d{4})', s)
    if m:
        mon, day, year = _MONTHS.get(m.group(1).lower()), int(m.group(2)), int(m.group(3))
        if mon:
            try:
                return datetime(year, mon, day)
            except ValueError:
                pass
    return None


def _quick_extract_metadata(text: str) -> dict:
    """Extract start_date, end_date, and value from document text using regex — no AI needed."""
    result: dict = {}
    lines = text[:8000]  # only scan the first ~8000 chars

    # ── Dates ──────────────────────────────────────────────────────────────────
    date_re = (
        r'(\d{4}-\d{2}-\d{2}'
        r'|\d{1,2}[/-]\d{1,2}[/-]\d{4}'
        r'|\d{1,2}\s+(?:January|February|March|April|May|June|July|August'
        r'|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun'
        r'|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+\d{4}'
        r'|(?:January|February|March|April|May|June|July|August'
        r'|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun'
        r'|Jul|Aug|Sep|Oct|Nov|Dec)\.?\s+\d{1,2},?\s+\d{4})'
    )
    start_kws = r'(?:effective\s+date|start\s+date|commencement\s+date|contract\s+date|date\s+of\s+agreement|agreement\s+date|signed.*?date|execution\s+date)'
    end_kws   = r'(?:end\s+date|expiry\s+date|expiration\s+date|termination\s+date|term\s+end|expires?\s+on|valid\s+(?:until|through|to))'

    for pattern, field in [
        (rf'(?i){start_kws}[:\s]+({date_re[1:-1]})', "start_date"),
        (rf'(?i){end_kws}[:\s]+({date_re[1:-1]})',   "end_date"),
    ]:
        m = re.search(pattern, lines)
        if m:
            d = _parse_date_str(m.group(1))
            if d:
                result[field] = d

    # Fallback: grab first two standalone dates found if still missing
    if not result.get("start_date") or not result.get("end_date"):
        all_dates = []
        for m in re.finditer(date_re, lines, re.IGNORECASE):
            d = _parse_date_str(m.group(0))
            if d and d not in all_dates:
                all_dates.append(d)
        if all_dates and not result.get("start_date"):
            result["start_date"] = all_dates[0]
        if len(all_dates) > 1 and not result.get("end_date"):
            result["end_date"] = all_dates[-1]

    # ── Value ──────────────────────────────────────────────────────────────────
    # Pass 1: "Total Contract Value" label — allow optional currency symbol between label and digits
    _CURRENCY = r'(?:USD|US\$|GBP|EUR|SGD|AUD|CAD|INR|RM|S\$|\$|£|€)?'
    val_m = re.search(
        r'(?i)(?:total\s+(?:contract\s+)?(?:value|amount|fee|price|cost)'
        r'|contract\s+(?:value|amount|price)'
        r'|aggregate\s+(?:contract\s+)?(?:value|amount))'
        r'[:\s]*' + _CURRENCY + r'\s*([\d,]+(?:\.\d{1,2})?)',
        text,  # search full text, not just first 8000 chars
    )
    # Pass 2: bare currency symbol + amount (first occurrence anywhere)
    if not val_m:
        val_m = re.search(
            r'(?:USD|US\$|GBP|EUR|SGD|AUD|CAD|\$|£|€)\s*([\d,]+(?:\.\d{1,2})?)',
            text,
            re.IGNORECASE,
        )
    if val_m:
        try:
            result["value"] = float(val_m.group(1).replace(",", ""))
        except ValueError:
            pass

    return result


@router.post("/", response_model=None)
async def create_new_contract(
    contract: ContractCreate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user_with_role),
):
    result = await create_contract(
        contract,
        user_id=current_user["user_id"],
        created_by_name=current_user.get("full_name") or current_user.get("email") or current_user["user_id"],
    )
    create_audit_log(
        action=AuditAction.create,
        resource_type="contract",
        resource_id=result.get("id", ""),
        user_id=current_user["user_id"],
        user_email=current_user.get("email"),
        details=f"Contract created: {contract.title}",
    )

    # Always generate an AI draft document in the background (blank or template).
    contracts_collection.update_one(
        {"_id": ObjectId(result["id"])},
        {"$set": {"document_status": "generating"}},
    )
    result["document_status"] = "generating"
    background_tasks.add_task(
        _ai_generate_document,
        contract_id=result["id"],
        contract_type=contract.contract_type.value,
        contract_title=contract.title,
        parties=[p.model_dump() for p in contract.parties],
        key_terms=_build_key_terms(contract),
        user_id=current_user["user_id"],
    )

    return result


def _build_key_terms(contract: ContractCreate) -> dict:
    """Build the key_terms dict passed to the AI draft generator."""
    terms: dict = {}

    # Standard contract metadata
    if contract.start_date:
        terms["start_date"] = contract.start_date.strftime("%d %B %Y")
        terms["effective_date"] = terms["start_date"]   # NDA template key
    if contract.end_date:
        terms["end_date"] = contract.end_date.strftime("%d %B %Y")
    if contract.value:
        terms["contract_value"] = str(contract.value)
    if contract.payment_terms:
        terms["payment_terms"] = contract.payment_terms
    if contract.description:
        terms["description"] = contract.description

    # Template-specific field values filled in by the user — these map directly
    # to {placeholders} in the AI template (e.g. purpose, governing_law, term…)
    if contract.template_values:
        terms.update(contract.template_values)

    return terms


async def _ai_generate_document(
    contract_id: str,
    contract_type: str,
    contract_title: str,
    parties: list,
    key_terms: dict,
    user_id: str,
):
    """Generate a contract document in two phases:
    Phase 1 (fast): build a template DOCX immediately so the user isn't kept waiting.
    Phase 2 (async): attempt AI enhancement; if successful, push an updated version.
    Risk analysis runs after Phase 2 regardless of which content was used.
    """
    import asyncio
    from bson import ObjectId as _OID
    from app.services.ai_service import generate_contract_draft, _mock_draft
    from app.services.document_export import html_to_docx

    safe_title = re.sub(r'[^\w\s-]', '', contract_title).strip().replace(" ", "_") or "contract"

    def _md_to_html(content: str) -> str:
        def _inline(text: str) -> str:
            text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<strong><em>\1</em></strong>', text)
            text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
            text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
            text = re.sub(r'_(.+?)_', r'<em>\1</em>', text)
            return text
        parts = []
        for line in content.split("\n"):
            s = line.strip()
            if s.startswith("### "):
                parts.append(f"<h3>{_inline(s[4:])}</h3>")
            elif s.startswith("## "):
                parts.append(f"<h2>{_inline(s[3:])}</h2>")
            elif s.startswith("# "):
                parts.append(f"<h1>{_inline(s[2:])}</h1>")
            elif s.startswith("- ") or s.startswith("* "):
                parts.append(f"<ul><li>{_inline(s[2:])}</li></ul>")
            elif re.match(r'^\d+\.\s+', s):
                item = re.sub(r'^\d+\.\s+', '', s)
                parts.append(f"<ol><li>{_inline(item)}</li></ol>")
            elif s in ("---", "***", "___"):
                parts.append("<hr>")
            elif s == "":
                parts.append("<br>")
            else:
                parts.append(f"<p>{_inline(s)}</p>")
        return "\n".join(parts)

    def _write_docx(content: str) -> tuple[str, int, bytes]:
        """Convert markdown content → DOCX, save to disk. Returns (stored_filename, size, bytes)."""
        html = _md_to_html(content)
        file_bytes = html_to_docx(html, title=contract_title)
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        fid = uuid.uuid4().hex
        fname = f"{fid}.docx"
        with open(os.path.join(UPLOAD_DIR, fname), "wb") as f:
            f.write(file_bytes)
        return fname, len(file_bytes), file_bytes

    # ── Phase 1: template DOCX — instant, marks document as ready ────────────
    try:
        template = _mock_draft(contract_type, parties)
        fname1, size1, bytes1 = _write_docx(template["content"])
        extracted_text = _extract_text_from_docx(bytes1)

        v1 = {
            "version_number": 1,
            "file_url": fname1,
            "original_filename": f"{safe_title}.docx",
            "file_size": size1,
            "file_type": ".docx",
            "uploaded_by": user_id,
            "uploaded_at": datetime.utcnow(),
            "change_notes": "Template draft",
        }
        set1 = {
            "file_url": fname1,
            "current_version": 1,
            "document_status": "ready",
            "updated_at": datetime.utcnow(),
        }
        if extracted_text:
            set1["extracted_text"] = extracted_text
        contracts_collection.update_one(
            {"_id": _OID(contract_id)},
            {"$push": {"versions": v1}, "$set": set1},
        )

    except Exception as exc:
        contracts_collection.update_one(
            {"_id": _OID(contract_id)},
            {"$set": {"document_status": "failed", "updated_at": datetime.utcnow()}},
        )
        return

    # ── Phase 2: AI enhancement — runs after user already has a document ─────
    try:
        ai_draft = await asyncio.wait_for(
            generate_contract_draft(contract_type=contract_type, parties=parties, key_terms=key_terms),
            timeout=45.0,
        )
        ai_content: str = ai_draft.get("content", "")
        # Only replace template if AI produced meaningfully more content
        if ai_content and len(ai_content) > 500:
            fname2, size2, bytes2 = _write_docx(ai_content)
            ai_text = _extract_text_from_docx(bytes2)
            v2 = {
                "version_number": 2,
                "file_url": fname2,
                "original_filename": f"{safe_title}.docx",
                "file_size": size2,
                "file_type": ".docx",
                "uploaded_by": user_id,
                "uploaded_at": datetime.utcnow(),
                "change_notes": "AI-generated draft",
            }
            set2 = {
                "file_url": fname2,
                "current_version": 2,
                "updated_at": datetime.utcnow(),
            }
            if ai_text:
                set2["extracted_text"] = ai_text
                extracted_text = ai_text
            contracts_collection.update_one(
                {"_id": _OID(contract_id)},
                {"$push": {"versions": v2}, "$set": set2},
            )
    except (asyncio.TimeoutError, Exception):
        pass

    # ── Risk analysis — runs regardless of which version was used ─────────────
    if extracted_text:
        try:
            from app.services.ai_service import analyze_contract_by_id
            await analyze_contract_by_id(contract_id)
        except Exception as exc:
            pass


@router.get("/")
async def list_contracts(
    search: Optional[str] = Query(None, description="Search by title"),
    contract_type: Optional[ContractType] = Query(None),
    status: Optional[ContractStatus] = Query(None),
    workflow_stage: Optional[WorkflowStage] = Query(None),
    risk_level: Optional[RiskLevel] = Query(None),
    page: int = Query(1, ge=1, description="Page number"),
    per_page: int = Query(20, ge=1, le=500, description="Items per page"),
    view: str = Query("active", description="active | archived | trash"),
    current_user: dict = Depends(get_current_user_with_role),
):
    is_admin = current_user.get("role") in ("admin", "manager")
    filters = ContractFilter(
        search=search,
        contract_type=contract_type,
        status=status,
        workflow_stage=workflow_stage,
        risk_level=risk_level,
        page=page,
        per_page=per_page,
    )
    return await get_contracts(filters, user_id=current_user["user_id"], is_admin=is_admin, view=view)


# Must be above /{contract_id} or FastAPI matches "dashboard" as an ID
@router.get("/dashboard")
async def dashboard_statistics():
    return await get_dashboard_stats()


# Must be above /{contract_id} or FastAPI matches "upload" as an ID
def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "\n\n".join(page.get_text() for page in doc)
    except Exception:
        return ""


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception:
        return ""


def _convert_pdf_to_docx(pdf_content: bytes) -> bytes | None:
    """Convert PDF bytes to DOCX using pdf2docx (pure Python, no external service needed)."""
    try:
        import tempfile
        from pdf2docx import Converter
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_tmp:
            pdf_tmp.write(pdf_content)
            pdf_tmp_path = pdf_tmp.name
        docx_tmp_path = pdf_tmp_path.replace(".pdf", ".docx")
        cv = Converter(pdf_tmp_path)
        cv.convert(docx_tmp_path, start=0, end=None)
        cv.close()
        with open(docx_tmp_path, "rb") as f:
            result = f.read()
        os.unlink(pdf_tmp_path)
        os.unlink(docx_tmp_path)
        return result
    except Exception:
        return None


@router.post("/upload")
async def upload_and_create_contract(
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = None,
    current_user: dict = Depends(get_current_user_with_role),
):
    """Upload a document (PDF/DOCX/TXT) and create a new draft contract from it."""
    # Validate extension
    original_name = file.filename or "untitled"
    _, ext = os.path.splitext(original_name)
    ext = ext.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{ext}' not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Read and validate size
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File exceeds 20 MB limit")

    # Save to disk
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    file_id = uuid.uuid4().hex
    stored_filename = f"{file_id}{ext}"
    file_path = os.path.join(UPLOAD_DIR, stored_filename)
    with open(file_path, "wb") as f:
        f.write(content)

    # For PDFs: convert to DOCX via pdf2docx so the document is fully editable.
    # The original PDF is kept as version 1; the DOCX becomes version 2 (working copy).
    pdf_original_stored = None
    if ext == ".pdf":
        docx_bytes = _convert_pdf_to_docx(content)
        if docx_bytes:
            pdf_original_stored = stored_filename
            docx_id = uuid.uuid4().hex
            stored_filename = f"{docx_id}.docx"
            docx_path = os.path.join(UPLOAD_DIR, stored_filename)
            with open(docx_path, "wb") as f:
                f.write(docx_bytes)
            content = docx_bytes
            ext = ".docx"
            original_name = os.path.splitext(original_name)[0] + ".docx"

    # Extract text for downstream AI analysis / display
    extracted_text = ""
    if pdf_original_stored:
        # PDF was converted to DOCX — extract from the original PDF
        extracted_text = _extract_text_from_pdf(open(os.path.join(UPLOAD_DIR, pdf_original_stored), "rb").read())
    elif ext == ".docx":
        extracted_text = _extract_text_from_docx(content)
    elif ext == ".pdf":
        # PDF-to-DOCX conversion failed — fall back to direct text extraction
        extracted_text = _extract_text_from_pdf(content)
    elif ext == ".txt":
        try:
            extracted_text = content.decode("utf-8", errors="replace")
        except Exception:
            extracted_text = ""

    # Derive a title from the filename (strip extension, replace underscores)
    title = os.path.splitext(original_name)[0].replace("_", " ").replace("-", " ").strip()
    if not title:
        title = "Uploaded Contract"

    now = datetime.utcnow()
    user_id = current_user["user_id"]

    version_entry = {
        "version_number": 1,
        "file_url": stored_filename,
        "original_filename": original_name,
        "file_size": len(content),
        "file_type": ext,
        "uploaded_by": user_id,
        "uploaded_at": now,
        "change_notes": "Initial upload",
    }

    contract_doc = {
        "title": title,
        "contract_type": "other",
        "description": f"Created from uploaded file: {original_name}",
        "parties": [],
        "start_date": None,
        "end_date": None,
        "value": None,
        "payment_terms": None,
        "status": "draft",
        "workflow_stage": "request",
        "approval_type": "all_required",
        "workflow_trigger": "creation",
        "file_url": stored_filename,
        "versions": [version_entry],
        "current_version": 1,
        "ai_analysis": None,
        "created_by": user_id,
        "created_by_name": current_user.get("full_name") or current_user.get("email") or user_id,
        "organization_id": None,
        "tags": ["uploaded"],
        "template_id": None,
        "created_at": now,
        "updated_at": now,
    }

    if extracted_text:
        contract_doc["extracted_text"] = extracted_text
        # Fast regex extraction — populates start_date, end_date, value immediately
        # without waiting for AI. AI analysis runs in background for risk/summary.
        quick = _quick_extract_metadata(extracted_text)
        for key in ("start_date", "end_date", "value"):
            if quick.get(key) is not None:
                contract_doc[key] = quick[key]

    result = contracts_collection.insert_one(contract_doc)
    contract_doc["id"] = str(result.inserted_id)
    del contract_doc["_id"]

    # Background AI analysis for risk score, summary, recommendations
    if extracted_text and background_tasks is not None:
        from app.services.ai_service import analyze_contract_by_id
        background_tasks.add_task(analyze_contract_by_id, contract_doc["id"])

    create_audit_log(
        action=AuditAction.file_upload,
        resource_type="contract",
        resource_id=contract_doc["id"],
        user_id=user_id,
        user_email=current_user.get("email"),
        details=f"Document uploaded: {original_name}",
    )

    return {
        "id": contract_doc["id"],
        "contract": contract_doc,
        "message": "Contract created from uploaded document",
        "extracted_text": extracted_text[:2000] if extracted_text else "",
    }


@router.get("/{contract_id}")
async def get_single_contract(
    contract_id: str,
    current_user: dict = Depends(get_current_user_with_role),
):
    is_admin = current_user.get("role") in ("admin", "manager")
    contract = await get_contract(contract_id, user_id=current_user["user_id"], is_admin=is_admin)

    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    return contract


@router.put("/{contract_id}")
async def update_existing_contract(
    contract_id: str,
    update_data: ContractUpdate,
    current_user: dict = Depends(get_current_user_with_role),
):
    is_admin = current_user.get("role") in ("admin", "manager")
    existing = await get_contract(contract_id, user_id=current_user["user_id"], is_admin=is_admin)
    if not existing:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract = await update_contract(contract_id, update_data)
    create_audit_log(
        action=AuditAction.update,
        resource_type="contract",
        resource_id=contract_id,
        user_id=current_user["user_id"],
        user_email=current_user.get("email"),
        details=f"Contract updated: {existing.get('title', contract_id)}",
    )
    return contract


@router.delete("/{contract_id}")
async def delete_existing_contract(
    contract_id: str,
    current_user: dict = Depends(get_current_user_with_role),
):
    """Soft-delete: moves contract to trash for 30 days before permanent removal."""
    is_admin = current_user.get("role") in ("admin", "manager")
    existing = await get_contract(contract_id, user_id=current_user["user_id"], is_admin=is_admin)
    if not existing:
        raise HTTPException(status_code=404, detail="Contract not found")

    await delete_contract(contract_id)
    create_audit_log(
        action=AuditAction.delete,
        resource_type="contract",
        resource_id=contract_id,
        user_id=current_user["user_id"],
        user_email=current_user.get("email"),
        details=f"Contract moved to trash: {existing.get('title', contract_id)}",
    )
    return {"message": "Contract moved to trash"}


@router.delete("/{contract_id}/permanent")
async def permanent_delete_existing_contract(
    contract_id: str,
    current_user: dict = Depends(get_current_user_with_role),
):
    """Permanently delete a contract from trash (irreversible)."""
    is_admin = current_user.get("role") in ("admin", "manager")
    if not is_admin:
        raise HTTPException(status_code=403, detail="Admin access required")

    deleted = await permanent_delete_contract(contract_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Contract not found")

    create_audit_log(
        action=AuditAction.delete,
        resource_type="contract",
        resource_id=contract_id,
        user_id=current_user["user_id"],
        user_email=current_user.get("email"),
        details=f"Contract permanently deleted: {contract_id}",
    )
    return {"message": "Contract permanently deleted"}


@router.patch("/{contract_id}/restore")
async def restore_contract_from_trash(
    contract_id: str,
    current_user: dict = Depends(get_current_user_with_role),
):
    """Restore a contract from trash back to active."""
    contract = await restore_contract(contract_id)
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found in trash")

    create_audit_log(
        action=AuditAction.update,
        resource_type="contract",
        resource_id=contract_id,
        user_id=current_user["user_id"],
        user_email=current_user.get("email"),
        details=f"Contract restored from trash: {contract.get('title', contract_id)}",
    )
    return contract


@router.patch("/{contract_id}/archive")
async def archive_existing_contract(
    contract_id: str,
    current_user: dict = Depends(get_current_user_with_role),
):
    """Archive a contract (hides it from the main list without deleting)."""
    is_admin = current_user.get("role") in ("admin", "manager")
    existing = await get_contract(contract_id, user_id=current_user["user_id"], is_admin=is_admin)
    if not existing:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract = await archive_contract(contract_id)
    create_audit_log(
        action=AuditAction.update,
        resource_type="contract",
        resource_id=contract_id,
        user_id=current_user["user_id"],
        user_email=current_user.get("email"),
        details=f"Contract archived: {existing.get('title', contract_id)}",
    )
    return contract


@router.patch("/{contract_id}/unarchive")
async def unarchive_existing_contract(
    contract_id: str,
    current_user: dict = Depends(get_current_user_with_role),
):
    """Move an archived contract back to the active list."""
    contract = await unarchive_contract(contract_id)
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")

    create_audit_log(
        action=AuditAction.update,
        resource_type="contract",
        resource_id=contract_id,
        user_id=current_user["user_id"],
        user_email=current_user.get("email"),
        details=f"Contract unarchived: {contract.get('title', contract_id)}",
    )
    return contract


@router.post("/purge-trash")
async def purge_old_trash(
    current_user: dict = Depends(get_current_user_with_role),
):
    """Permanently delete all contracts that have been in trash for more than 30 days."""
    is_admin = current_user.get("role") in ("admin", "manager")
    if not is_admin:
        raise HTTPException(status_code=403, detail="Admin access required")

    count = await purge_expired_trash()
    return {"message": f"Purged {count} expired contracts from trash"}


@router.patch("/{contract_id}/workflow")
async def change_workflow_stage(
    contract_id: str,
    stage: WorkflowStage,
    current_user: dict = Depends(get_current_user_with_role),
):
    is_admin = current_user.get("role") in ("admin", "manager")
    existing = await get_contract(contract_id, user_id=current_user["user_id"], is_admin=is_admin)
    if not existing:
        raise HTTPException(status_code=404, detail="Contract not found")

    contract = await update_workflow_stage(contract_id, stage.value)
    create_audit_log(
        action=AuditAction.status_change,
        resource_type="contract",
        resource_id=contract_id,
        user_id=current_user["user_id"],
        user_email=current_user.get("email"),
        details=f"Workflow stage changed to: {stage.value} on contract: {existing.get('title', contract_id)}",
    )
    return contract


class AddVersionBody(BaseModel):
    from_contract_id: str
    change_notes: str = "New version"


@router.post("/{contract_id}/add-version")
async def add_version_to_contract(
    contract_id: str,
    body: AddVersionBody,
    current_user: dict = Depends(get_current_user_with_role),
):
    """Absorb a newly uploaded contract as a new version of an existing one, then delete the source."""
    if not ObjectId.is_valid(contract_id) or not ObjectId.is_valid(body.from_contract_id):
        raise HTTPException(status_code=400, detail="Invalid contract ID")

    target = contracts_collection.find_one({"_id": ObjectId(contract_id)})
    if not target:
        raise HTTPException(status_code=404, detail="Target contract not found")

    source = contracts_collection.find_one({"_id": ObjectId(body.from_contract_id)})
    if not source:
        raise HTTPException(status_code=404, detail="Source contract not found")

    new_version_number = (target.get("current_version") or 1) + 1
    source_v1 = (source.get("versions") or [{}])[0]

    new_version_entry = {
        "version_number": new_version_number,
        "file_url": source.get("file_url", ""),
        "original_filename": source_v1.get("original_filename", source.get("file_url", "")),
        "file_size": source_v1.get("file_size", 0),
        "file_type": source_v1.get("file_type", ""),
        "uploaded_by": current_user["user_id"],
        "uploaded_at": datetime.utcnow(),
        "change_notes": body.change_notes,
    }

    set_fields = {
        "current_version": new_version_number,
        "file_url": source.get("file_url", ""),
        "updated_at": datetime.utcnow(),
    }
    if source.get("extracted_text"):
        set_fields["extracted_text"] = source["extracted_text"]

    contracts_collection.update_one(
        {"_id": ObjectId(contract_id)},
        {"$push": {"versions": new_version_entry}, "$set": set_fields},
    )
    contracts_collection.delete_one({"_id": ObjectId(body.from_contract_id)})

    create_audit_log(
        action=AuditAction.update,
        resource_type="contract",
        resource_id=contract_id,
        user_id=current_user["user_id"],
        user_email=current_user.get("email"),
        details=f"Version {new_version_number} added ({body.change_notes})",
    )
    return {
        "message": f"Added as version {new_version_number}",
        "version_number": new_version_number,
        "contract_id": contract_id,
    }